from __future__ import annotations

from io import BytesIO

from docx import Document
from openpyxl import load_workbook
from pptx import Presentation
from pypdf import PdfReader

from app.models import Artifact
from app.services.artifacts import ArtifactService


def test_all_document_generators_create_valid_files(app):
    sources = [
        {
            "title": "Primary source",
            "url": "https://example.com/evidence",
            "retrieved_at": "2026-07-24T00:00:00+00:00",
        }
    ]
    with app.app_context():
        service = ArtifactService("artifact-owner")
        excel = service.create_excel(
            "Monthly Sales",
            [
                {"Month": "January", "Sales": 12000, "Growth": 0.12},
                {"Month": "February", "Sales": 14800, "Growth": 0.23},
            ],
            sources=sources,
        )
        word = service.create_word(
            "Sales Review",
            [
                {"heading": "Executive summary", "body": "Sales increased during the period."},
                {
                    "heading": "Results",
                    "body": ["Revenue grew.", "Conversion improved."],
                    "table": [{"Metric": "Sales", "Value": "14,800"}],
                },
            ],
            sources=sources,
        )
        slides = service.create_powerpoint(
            "Sales Review",
            [
                {
                    "title": "Momentum",
                    "bullets": ["Sales increased in February", "Growth reached 23%"],
                    "takeaway": "The strongest month was February.",
                },
                {
                    "title": "Next steps",
                    "bullets": ["Validate pipeline quality", "Protect gross margin"],
                    "takeaway": "Scale the repeatable channels.",
                },
            ],
            sources=sources,
        )
        pdf = service.create_pdf(
            "Sales Review",
            [{"heading": "Executive summary", "body": "Sales increased during the period."}],
            sources=sources,
        )

        workbook = load_workbook(excel.storage_path, data_only=False)
        assert workbook.sheetnames == ["Data", "Charts", "Metadata"]
        assert workbook["Data"].freeze_panes == "A5"
        assert workbook["Data"].tables
        assert workbook["Metadata"]["B8"].value == "https://example.com/evidence"
        workbook.close()

        document = Document(word.storage_path)
        assert document.core_properties.title == "Sales Review"
        assert any(paragraph.text == "References" for paragraph in document.paragraphs)

        presentation = Presentation(slides.storage_path)
        assert presentation.slide_width / presentation.slide_height > 1.7
        assert len(presentation.slides) == 5

        reader = PdfReader(pdf.storage_path)
        assert len(reader.pages) >= 2
        assert "Sales Review" in (reader.pages[0].extract_text() or "")
        assert Artifact.query.filter_by(owner_id="artifact-owner").count() == 4


def test_excel_neutralizes_formula_injection(app):
    with app.app_context():
        artifact = ArtifactService("formula-owner").create_excel(
            "Unsafe CSV values",
            [{"Name": '=HYPERLINK("https://evil.example")', "Value": 12}],
        )
        workbook = load_workbook(artifact.storage_path, data_only=False)
        assert workbook["Data"]["A5"].value.startswith("'=")
        workbook.close()


def test_generated_excel_download_is_valid(client, csrf_headers):
    conversation = client.post("/api/conversations", json={}, headers=csrf_headers).get_json()
    upload = client.post(
        "/api/uploads",
        data={"file": (BytesIO(b"Month,Sales\nJan,120\nFeb,180\n"), "sales.csv")},
        headers=csrf_headers,
        content_type="multipart/form-data",
    )
    assert upload.status_code == 201
    upload_id = upload.get_json()["id"]
    plan_response = client.post(
        "/api/plans",
        json={
            "conversation_id": conversation["id"],
            "goal": "Create an Excel workbook from my uploaded sales CSV with charts and insights.",
            "attachment_ids": [upload_id],
        },
        headers=csrf_headers,
    )
    assert plan_response.status_code == 201
    plan_id = plan_response.get_json()["plan"]["id"]
    execution = client.post(f"/api/plans/{plan_id}/execute", json={}, headers=csrf_headers)
    assert execution.status_code == 200
    assert b"event: done" in execution.data

    artifacts = client.get("/api/artifacts").get_json()
    assert len(artifacts) == 1
    download = client.get(artifacts[0]["download_url"])
    assert download.status_code == 200
    workbook = load_workbook(BytesIO(download.data), read_only=True)
    assert "Data" in workbook.sheetnames
    assert "Charts" in workbook.sheetnames
    workbook.close()
